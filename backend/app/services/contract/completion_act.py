"""Генерация DOCX акта выполненных работ (об оказании услуг)."""

from __future__ import annotations

import tempfile
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from app.core.config import settings

from .docx_utils import _cell_add_break_lines
from .number_format import fmt_rub, number_to_words_ru, ru_date


def generate_completion_act(
    order_id_short: str,
    contract_number: str,
    object_address: str,
    payment_amount: int,
    requisites: dict[str, Any],
) -> Path:
    """Генерирует DOCX акта выполненных работ (об оказании услуг).

    Args:
        order_id_short: Первые 8 символов UUID заявки (для имени файла).
        contract_number: Номер договора.
        object_address:  Адрес объекта.
        payment_amount:  Полная стоимость работ, руб.
        requisites:      Реквизиты клиента из CompanyRequisites.model_dump().

    Returns:
        Path к временному .docx файлу. Вызывающий код удаляет после отправки.
    """
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(2)
    sec.right_margin = Cm(1.5)

    s = settings
    now = datetime.now()
    client_kpp = requisites.get("kpp") or "—"
    words_amount = number_to_words_ru(payment_amount)

    # ── Шапка: реквизиты банка получателя ────────────────────────────────────
    hdr = doc.add_table(rows=4, cols=2)
    hdr.style = "Table Grid"
    hdr.autofit = False
    hdr.columns[0].width = Cm(10)
    hdr.columns[1].width = Cm(7.5)

    def _hdr(row: int, col: int, lines: str, bold: bool = False) -> None:
        cell = hdr.rows[row].cells[col]
        cell.paragraphs[0].clear()
        _cell_add_break_lines(cell, lines)
        if bold:
            for run in cell.paragraphs[0].runs:
                run.bold = True

    _hdr(0, 0, f"Банк получателя:\n{s.company_bank_name}")
    _hdr(0, 1, f"БИК\n{s.company_bik}")
    _hdr(1, 0, "")
    _hdr(1, 1, f"Сч. №\n{s.company_corr_account}")
    _hdr(2, 0, f"ИНН {s.company_inn}")
    _hdr(2, 1, "")
    _hdr(3, 0, f"Получатель:\n{s.company_full_name}", bold=True)
    _hdr(3, 1, f"Сч. №\n{s.company_settlement_account}")

    doc.add_paragraph()

    # ── Заголовок акта ────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        f"АКТ ОБ ОКАЗАНИИ УСЛУГ № {contract_number}\nот {ru_date(now)}"
    )
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()

    # ── Поставщик и покупатель ────────────────────────────────────────────────
    def _info_line(label: str, value: str) -> None:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r1 = p.add_run(f"{label}: ")
        r1.bold = True
        r1.font.size = Pt(11)
        r2 = p.add_run(value)
        r2.font.size = Pt(11)

    _info_line(
        "Исполнитель",
        f"{s.company_full_name}, ИНН {s.company_inn}, {s.company_address or '—'}",
    )
    _info_line(
        "Заказчик",
        f"{requisites['full_name']}, ИНН {requisites['inn']}, "
        f"КПП {client_kpp}, {requisites['legal_address']}",
    )
    _info_line("Основание", f"Договор № {contract_number}")

    doc.add_paragraph()

    # ── Таблица услуг ─────────────────────────────────────────────────────────
    col_headers = ["№", "Наименование работ (услуг)", "Кол.", "Ед.", "Цена, руб.", "Сумма, руб."]
    col_widths = [Cm(1.0), Cm(8.5), Cm(1.5), Cm(1.5), Cm(2.5), Cm(2.5)]

    svc_tbl = doc.add_table(rows=2, cols=6)
    svc_tbl.style = "Table Grid"
    svc_tbl.autofit = False

    for i, (header, width) in enumerate(zip(col_headers, col_widths)):
        cell = svc_tbl.rows[0].cells[i]
        cell.width = width
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(header)
        run.bold = True
        run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    service_name = f"Разработка проекта УУТЭ по объекту:\n{object_address}"
    row_data = [
        ("1", WD_ALIGN_PARAGRAPH.CENTER),
        (service_name, WD_ALIGN_PARAGRAPH.LEFT),
        ("1", WD_ALIGN_PARAGRAPH.CENTER),
        ("усл.", WD_ALIGN_PARAGRAPH.CENTER),
        (fmt_rub(payment_amount), WD_ALIGN_PARAGRAPH.RIGHT),
        (fmt_rub(payment_amount), WD_ALIGN_PARAGRAPH.RIGHT),
    ]
    for i, (text, align) in enumerate(row_data):
        cell = svc_tbl.rows[1].cells[i]
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(text)
        run.font.size = Pt(10)
        cell.paragraphs[0].alignment = align

    doc.add_paragraph()

    # ── Итого ─────────────────────────────────────────────────────────────────
    for text, bold in [
        (f"Итого: {fmt_rub(payment_amount)} руб.", True),
        ("В том числе НДС: не облагается (УСН, ст. 346.11 НК РФ).", False),
        (
            f"Всего по акту: {fmt_rub(payment_amount)} руб. "
            f"({words_amount[0].upper() + words_amount[1:]}).",
            True,
        ),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(12)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(
        "Вышеперечисленные услуги выполнены в полном объёме, "
        "в установленные сроки, стороны претензий друг к другу не имеют."
    )
    run.font.size = Pt(11)

    doc.add_paragraph()
    doc.add_paragraph()

    # ── Подписи двух сторон ───────────────────────────────────────────────────
    sig_tbl = doc.add_table(rows=3, cols=2)
    sig_tbl.autofit = False
    sig_tbl.columns[0].width = Cm(8.75)
    sig_tbl.columns[1].width = Cm(8.75)

    def _sig_cell(row: int, col: int, text: str) -> None:
        cell = sig_tbl.rows[row].cells[col]
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(text)
        run.font.size = Pt(11)
        if row == 0:
            run.bold = True

    _sig_cell(0, 0, "ИСПОЛНИТЕЛЬ")
    _sig_cell(0, 1, "ЗАКАЗЧИК")
    _sig_cell(1, 0, s.company_full_name)
    _sig_cell(1, 1, requisites["full_name"])
    _sig_cell(
        2, 0,
        f"{s.company_director_position}: __________ / {s.company_director_name}",
    )
    director_pos = requisites.get("director_position") or "Генеральный директор"
    director_name = requisites.get("director_name") or "________________"
    _sig_cell(2, 1, f"{director_pos}: __________ / {director_name}")

    # ── Сохранить во временный файл ───────────────────────────────────────────
    tmp = tempfile.NamedTemporaryFile(
        suffix=".docx",
        prefix=f"akt_{order_id_short}_",
        delete=False,
    )
    tmp.close()
    doc.save(tmp.name)
    return Path(tmp.name)
