"""
* @file: docx_utils.py
* @description: Общие утилиты python-docx — таблицы, параграфы, отступы договора.
* @created: 2026-04-22
"""

from __future__ import annotations

from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

_CONTRACT_FONT_SIZE_PT = 10


def _compact_paragraph(paragraph: Any, *, size: int = _CONTRACT_FONT_SIZE_PT) -> None:
    """Делает параграф компактным для плотного договора без лишних интервалов."""
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.paragraph_format.line_spacing = 1.0
    for run in paragraph.runs:
        run.font.size = Pt(size)


def _set_table_borders(
    table: Any,
    *,
    outer: bool = False,
    inner_h: bool = False,
    inner_v: bool = False,
) -> None:
    """Управляет видимостью границ таблицы через XML."""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)

    borders_el = OxmlElement("w:tblBorders")
    nil = {"val": "nil"}
    single = {"val": "single", "sz": "6", "space": "0", "color": "auto"}

    mapping = {
        "top": single if outer else nil,
        "left": single if outer else nil,
        "bottom": single if outer else nil,
        "right": single if outer else nil,
        "insideH": single if inner_h else nil,
        "insideV": single if inner_v else nil,
    }
    for edge, attrs in mapping.items():
        el = OxmlElement(f"w:{edge}")
        for k, v in attrs.items():
            el.set(qn(f"w:{k}"), v)
        borders_el.append(el)
    tblPr.append(borders_el)


def _cell_lines(cell: Any, lines: list[Any], font_size: int = 10) -> None:
    """Заполняет ячейку несколькими строками (строка или кортеж (строка, bold))."""
    for i, line in enumerate(lines):
        text, bold = line if isinstance(line, tuple) else (line, False)
        para = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        _compact_paragraph(para, size=font_size)
        run = para.add_run(text)
        run.font.size = Pt(font_size)
        run.bold = bold


def _cell_add_break_lines(cell: Any, text: str, font_size: int = 9) -> None:
    """Добавляет текст с переносами строк (\\n) в ячейку через w:br."""
    para = cell.paragraphs[0]
    lines = text.split("\n")
    for i, line in enumerate(lines):
        if i > 0:
            if para.runs:
                para.runs[-1].add_break()
            else:
                para.add_run().add_break()
        run = para.add_run(line)
        run.font.size = Pt(font_size)


def _apply_contract_page_settings(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(3)
    sec.right_margin = Cm(1.5)
    doc.styles["Normal"].font.size = Pt(_CONTRACT_FONT_SIZE_PT)


def _add_paragraph(
    doc: Document,
    text: str,
    *,
    bold: bool = False,
    size: int = _CONTRACT_FONT_SIZE_PT,
    align: Any = WD_ALIGN_PARAGRAPH.JUSTIFY,
) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    _compact_paragraph(p, size=size)


def _section_heading(doc: Document, title: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(_CONTRACT_FONT_SIZE_PT)
    _compact_paragraph(p)


def _add_justify_lines(doc: Document, lines: list[str]) -> None:
    for text in lines:
        _add_paragraph(doc, text)
