"""Генерация DOCX счёта (аванс / остаток)."""

from __future__ import annotations

import tempfile
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from app.core.config import settings

from .docx_utils import _cell_add_break_lines
from .number_format import fmt_rub, number_to_words_ru, ru_date


def generate_invoice(
    order_id_short: str,
    contract_number: str,
    object_address: str,
    payment_amount: int,
    advance_amount: int,
    requisites: dict[str, Any],
    is_advance: bool = True,
) -> Path:
    """Генерирует DOCX счёта на оплату (аванс или окончательный расчёт).

    Args:
        order_id_short: Первые 8 символов UUID заявки (для имени файла).
        contract_number: Номер договора.
        object_address:  Адрес объекта.
        payment_amount:  Полная стоимость работ, руб.
        advance_amount:  Сумма аванса (50%), руб.
        requisites:      Реквизиты клиента из CompanyRequisites.model_dump().
        is_advance:      True = счёт на 50% аванс, False = счёт на остаток 50%.

    Returns:
        Path к временному .docx файлу. Вызывающий код удаляет после отправки.
    """
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(2)
    sec.right_margin = Cm(1.5)

    s = settings
    amount = advance_amount if is_advance else (payment_amount - advance_amount)
    invoice_label = "аванс 50%" if is_advance else "окончательный расчёт 50%"
    now = datetime.now()
    client_kpp = requisites.get("kpp") or "—"
    words_amount = number_to_words_ru(amount)

    # ── Шапка: реквизиты банка получателя ────────────────────────────────────
    # Стандартный формат российского счёта: 4 строки, 2 колонки
    hdr = doc.add_table(rows=4, cols=2)
    hdr.style = "Table Grid"
    hdr.autofit = False
    hdr.columns[0].width = Cm(10)
    hdr.columns[1].width = Cm(7.5)

    def _hdr(row: int, col: int, lines: str, bold: bool = False) -> None:
        cell = hdr.rows[row].cells[col]
        cell.paragraphs[0].clear()
        _cell_add_break_lines(cell, lines)
        if bold:
            for run in cell.paragraphs[0].runs:
                run.bold = True

    _hdr(0, 0, f"Банк получателя:\n{s.company_bank_name}")
    _hdr(0, 1, f"БИК\n{s.company_bik}")
    _hdr(1, 0, "")
    _hdr(1, 1, f"Сч.\u202f№\n{s.company_corr_account}")
    _hdr(2, 0, f"ИНН\u202f{s.company_inn}")
    _hdr(2, 1, "")
    _hdr(3, 0, f"Получатель:\n{s.company_full_name}", bold=True)
    _hdr(3, 1, f"Сч.\u202f№\n{s.company_settlement_account}")

    doc.add_paragraph()

    # ── Заголовок счёта ───────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"СЧЁТ НА ОПЛАТУ\u202f№\u202f{contract_number}\nот {ru_date(now)}")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()

    # ── Поставщик и покупатель ────────────────────────────────────────────────
    def _info_line(label: str, value: str) -> None:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r1 = p.add_run(f"{label}:\u2003")
        r1.bold = True
        r1.font.size = Pt(11)
        r2 = p.add_run(value)
        r2.font.size = Pt(11)

    _info_line(
        "Поставщик (Исполнитель)",
        f"{s.company_full_name}, ИНН\u202f{s.company_inn}, {s.company_address or '—'}",
    )
    _info_line(
        "Покупатель (Заказчик)",
        f"{requisites['full_name']}, ИНН\u202f{requisites['inn']}, "
        f"КПП\u202f{client_kpp}, {requisites['legal_address']}",
    )
    _info_line("Основание", f"Договор\u202f№\u202f{contract_number}")

    doc.add_paragraph()

    # ── Таблица услуг ─────────────────────────────────────────────────────────
    # Ширины: 1.0 + 8.5 + 1.5 + 1.5 + 2.5 + 2.5 = 17.5 см
    col_headers = ["№", "Наименование работ (услуг)", "Кол.", "Ед.", "Цена, руб.", "Сумма, руб."]
    col_widths = [Cm(1.0), Cm(8.5), Cm(1.5), Cm(1.5), Cm(2.5), Cm(2.5)]

    svc_tbl = doc.add_table(rows=2, cols=6)
    svc_tbl.style = "Table Grid"
    svc_tbl.autofit = False

    # Заголовок таблицы
    for i, (header, width) in enumerate(zip(col_headers, col_widths)):
        cell = svc_tbl.rows[0].cells[i]
        cell.width = width
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(header)
        run.bold = True
        run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Строка с услугой
    service_name = f"Разработка проекта УУТЭ по объекту:\n{object_address} ({invoice_label})"
    row_data = [
        ("1", WD_ALIGN_PARAGRAPH.CENTER),
        (service_name, WD_ALIGN_PARAGRAPH.LEFT),
        ("1", WD_ALIGN_PARAGRAPH.CENTER),
        ("усл.", WD_ALIGN_PARAGRAPH.CENTER),
        (fmt_rub(amount), WD_ALIGN_PARAGRAPH.RIGHT),
        (fmt_rub(amount), WD_ALIGN_PARAGRAPH.RIGHT),
    ]
    for i, (text, align) in enumerate(row_data):
        cell = svc_tbl.rows[1].cells[i]
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(text)
        run.font.size = Pt(10)
        cell.paragraphs[0].alignment = align

    doc.add_paragraph()

    # ── Итого ─────────────────────────────────────────────────────────────────
    for text, bold in [
        (f"Итого:\u2003{fmt_rub(amount)}\u202fруб.", True),
        ("В том числе НДС: не облагается (УСН, ст.\u202f346.11 НК\u202fРФ).", False),
        (
            f"Всего к оплате:\u2003{fmt_rub(amount)}\u202fруб. "
            f"({words_amount[0].upper() + words_amount[1:]}).",
            True,
        ),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph()

    # ── Подпись ───────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(
        f"{s.company_director_position}:\u2003__________\u202f/\u202f{s.company_director_name}"
    )
    run.font.size = Pt(12)

    # ── Сохранить во временный файл ───────────────────────────────────────────
    tmp = tempfile.NamedTemporaryFile(
        suffix=".docx",
        prefix=f"schet_{order_id_short}_",
        delete=False,
    )
    tmp.close()
    doc.save(tmp.name)
    return Path(tmp.name)
