"""Генератор договора на проектирование УУТЭ и счёта на аванс (.docx).

Использует python-docx (уже в requirements.txt) по аналогии с cover_letter.py.
Реквизиты исполнителя берутся из settings.company_*.

Приложение №2: страницы PDF ТУ (FileCategory.TU) встраиваются как растровые
изображения с автоматическим снижением DPI, чтобы итоговый DOCX укладывался
в лимит почтовых вложений (~25 МБ).

Функции:
- generate_contract_number(order_id) → str
- generate_contract(...)             → Path к временному .docx
- generate_invoice(...)              → Path к временному .docx

Вызывающий код обязан удалить временные файлы после отправки.
"""

import logging
import os
import re
import tempfile
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path

import fitz  # PyMuPDF

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from app.core.config import settings


logger = logging.getLogger(__name__)

# Лестница DPI для растра ТУ: сначала максимальное качество, затем снижение,
# пока размер DOCX не уложится в лимит для SMTP Яндекса (~30 МБ вложение).
_TU_DPI_LADDER = [150, 120, 100]
_MAX_DOCX_SIZE_BYTES = 25 * 1024 * 1024  # 25 МБ с запасом до Яндекс 30 МБ


# ═══════════════════════════════════════════════════════════════════════════════
# Вспомогательные утилиты
# ═══════════════════════════════════════════════════════════════════════════════

_RU_MONTHS = [
    "", "января", "февраля", "марта", "апреля", "мая", "июня",
    "июля", "августа", "сентября", "октября", "ноября", "декабря",
]

_ONES_M = ["", "один", "два", "три", "четыре", "пять",
           "шесть", "семь", "восемь", "девять"]
_ONES_F = ["", "одна", "две", "три", "четыре", "пять",
           "шесть", "семь", "восемь", "девять"]
_TEENS  = ["десять", "одиннадцать", "двенадцать", "тринадцать", "четырнадцать",
           "пятнадцать", "шестнадцать", "семнадцать", "восемнадцать", "девятнадцать"]
_TENS   = ["", "десять", "двадцать", "тридцать", "сорок", "пятьдесят",
           "шестьдесят", "семьдесят", "восемьдесят", "девяносто"]
_HUNDS  = ["", "сто", "двести", "триста", "четыреста", "пятьсот",
           "шестьсот", "семьсот", "восемьсот", "девятьсот"]


def _ru_date(dt: datetime) -> str:
    """Дата в русском формате: «15 апреля 2026 г.»"""
    return f"{dt.day}\u202f{_RU_MONTHS[dt.month]}\u202f{dt.year}\u202fг."


def _decline(n: int, one: str, few: str, many: str) -> str:
    """Склоняет слово по числу (рубль/рубля/рублей)."""
    if 11 <= n % 100 <= 19:
        return many
    m = n % 10
    if m == 1:
        return one
    if 2 <= m <= 4:
        return few
    return many


def _say_hundreds(n: int, feminine: bool = False) -> list[str]:
    """Произносит число 1–999 в виде списка слов."""
    ones = _ONES_F if feminine else _ONES_M
    words: list[str] = []
    h, r = divmod(n, 100)
    if h:
        words.append(_HUNDS[h])
    if 10 <= r <= 19:
        words.append(_TEENS[r - 10])
    else:
        t, u = divmod(r, 10)
        if t:
            words.append(_TENS[t])
        if u:
            words.append(ones[u])
    return words


def number_to_words_ru(n: int) -> str:
    """Число прописью (рубли), диапазон 0–999 999.

    Примеры:
        22500 → «двадцать две тысячи пятьсот рублей»
        11000 → «одиннадцать тысяч рублей»
        1     → «один рубль»
    """
    if n == 0:
        return "ноль рублей"
    parts: list[str] = []
    thousands, remainder = divmod(n, 1000)
    if thousands:
        parts.extend(_say_hundreds(thousands, feminine=True))
        parts.append(_decline(thousands, "тысяча", "тысячи", "тысяч"))
    if remainder:
        parts.extend(_say_hundreds(remainder, feminine=False))
    parts.append(_decline(n, "рубль", "рубля", "рублей"))
    return " ".join(parts)


def _fmt(n: int) -> str:
    """Форматирует сумму с узким неразрывным пробелом как разделителем тысяч."""
    return f"{n:,}".replace(",", "\u202f")


def _extract_city(address: str) -> str:
    """Извлекает город из адреса (первая часть до запятой, без «г. »)."""
    city_part = address.split(",")[0].strip()
    return re.sub(r"^г\.\s*", "", city_part).strip()


# ═══════════════════════════════════════════════════════════════════════════════
# XML-утилиты для таблиц
# ═══════════════════════════════════════════════════════════════════════════════


def _set_table_borders(
    table,
    *,
    outer: bool = False,
    inner_h: bool = False,
    inner_v: bool = False,
) -> None:
    """Управляет видимостью границ таблицы через XML.

    Параметры управляют включением отдельных групп границ:
      outer   — внешний контур (top, left, bottom, right)
      inner_h — горизонтальные разделители между строками
      inner_v — вертикальный разделитель между колонками
    """
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)

    borders_el = OxmlElement("w:tblBorders")
    nil = {"val": "nil"}
    single = {"val": "single", "sz": "6", "space": "0", "color": "auto"}

    mapping = {
        "top":     single if outer   else nil,
        "left":    single if outer   else nil,
        "bottom":  single if outer   else nil,
        "right":   single if outer   else nil,
        "insideH": single if inner_h else nil,
        "insideV": single if inner_v else nil,
    }
    for edge, attrs in mapping.items():
        el = OxmlElement(f"w:{edge}")
        for k, v in attrs.items():
            el.set(qn(f"w:{k}"), v)
        borders_el.append(el)
    tblPr.append(borders_el)


def _cell_lines(cell, lines: list, font_size: int = 10) -> None:
    """Заполняет ячейку несколькими строками.

    Каждый элемент — строка или кортеж (строка, bold).
    Первая строка использует существующий параграф ячейки,
    последующие — добавляются как новые параграфы.
    """
    for i, line in enumerate(lines):
        text, bold = line if isinstance(line, tuple) else (line, False)
        para = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        run = para.add_run(text)
        run.font.size = Pt(font_size)
        run.bold = bold


def _cell_add_break_lines(cell, text: str, font_size: int = 9) -> None:
    """Добавляет текст с переносами строк (\\n) в ячейку через w:br."""
    para = cell.paragraphs[0]
    lines = text.split("\n")
    for i, line in enumerate(lines):
        if i > 0:
            para.runs[-1].add_break() if para.runs else para.add_run().add_break()
        run = para.add_run(line)
        run.font.size = Pt(font_size)


# ═══════════════════════════════════════════════════════════════════════════════
# ТУ: растр PDF → PNG, сборка договора по шаблону kontrakt_ukute_template.md
# ═══════════════════════════════════════════════════════════════════════════════


def _render_tu_pages_to_png(
    tu_pdf_path: Path,
    dpi: int,
    out_dir: Path | None = None,
) -> list[Path]:
    """Рендерит каждую страницу PDF ТУ в PNG-файл на диске.

    Args:
        tu_pdf_path: путь к PDF ТУ
        dpi: разрешение рендера (150/120/100)
        out_dir: директория для PNG (по умолчанию — tempdir под `/tmp`)

    Returns:
        Список Path к PNG. Пустой список если файл не читается.
    """
    if out_dir is None:
        out_dir = Path(
            tempfile.mkdtemp(prefix="tu_pages_", dir="/tmp"),
        )

    paths: list[Path] = []
    try:
        pdf = fitz.open(str(tu_pdf_path))
    except Exception as exc:
        logger.warning("Не удалось открыть ТУ %s: %s", tu_pdf_path, exc)
        return []

    try:
        for page_num in range(len(pdf)):
            pix = pdf[page_num].get_pixmap(
                dpi=dpi,
            )
            png_path = out_dir / f"tu_p{page_num:03d}.png"
            pix.save(str(png_path))
            paths.append(png_path)
    finally:
        pdf.close()

    return paths


def _cleanup_png_files(paths: list[Path]) -> None:
    """Удаляет временные PNG-файлы, игнорирует ошибки.

    После удаления файлов пробует удалить родительский каталог (каталог ``tu_pages_*``
    из ``mkdtemp``), если он опустел; игнорирует сбои (непустая папка, race).
    """
    for p in paths:
        try:
            p.unlink()
        except OSError:
            pass
    # Удалить родительскую директорию если пустая
    if paths:
        try:
            paths[0].parent.rmdir()
        except OSError:
            pass


@dataclass
class _ContractContext:
    """Срез данных для заполнения DOCX договора (основной текст и приложения)."""

    order_id_short: str
    contract_number: str
    object_address: str
    client_name: str
    payment_amount: int
    advance_amount: int
    final_amount: int
    requisites: dict
    client_email: str | None
    rso_name: str | None
    tu_number: str | None
    tu_date: str | None
    tu_valid_to: str | None


def _apply_contract_page_settings(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(3)
    sec.right_margin = Cm(1.5)


def _add_paragraph(
    doc: Document,
    text: str,
    *,
    bold: bool = False,
    size: int = 12,
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold


def _section_heading(doc: Document, title: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(12)


def _add_justify_lines(doc: Document, lines: list[str]) -> None:
    for text in lines:
        _add_paragraph(doc, text)


def _fmt_word_money(n: int) -> str:
    w = number_to_words_ru(n)
    return w[0].upper() + w[1:] if w else w


def _appendix_signatures_row(doc: Document) -> None:
    """Подписи Исполнитель / Заказчик для конца приложения (как в шаблоне)."""
    s = settings
    req = _current_req_snapshot or {}
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(8)
    table.columns[1].width = Cm(8)
    _set_table_borders(table, inner_v=True)
    left_cell, right_cell = table.rows[0].cells
    _cell_lines(left_cell, [
        ("Исполнитель", True),
        "",
        f"__________\u202f/\u202f{s.company_director_name}",
        "М.П.",
    ])
    _cell_lines(right_cell, [
        ("Заказчик", True),
        "",
        f"__________\u202f/\u202f{req['director_name']}",
        "М.П.",
    ])


# Кэш requisites на время одного вызова сборки — для _appendix_signatures_row;
# задаётся внутри _build_full_contract_document.
_current_req_snapshot: dict | None = None


def _build_main_contract(doc: Document, ctx: _ContractContext) -> None:
    """Разделы 1–15 основного договора (по docs/kontrakt_ukute_template.md)."""
    s = settings
    city = _extract_city(s.company_address) if s.company_address else "___"
    now = datetime.now()
    req = ctx.requisites
    client_kpp = req.get("kpp") or "—"
    words_total = _fmt_word_money(ctx.payment_amount)
    words_adv = _fmt_word_money(ctx.advance_amount)
    words_fin = _fmt_word_money(ctx.final_amount)
    client_contact = ctx.client_email or req.get("email") or ctx.client_name
    executor_email = s.smtp_from or "—"
    ce = client_contact or "—"

    for title in [
        f"ДОГОВОР\u202f№\u202f{ctx.contract_number}",
        "на выполнение рабочей документации узла коммерческого учёта "
        "тепловой энергии, теплоносителя",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(14)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r1 = p.add_run(f"г.\u202f{city}")
    r1.font.size = Pt(12)
    r2 = p.add_run(
        f"\t\t\t«____»\u202f__________\u202f{now.year}\u202fг."
    )
    r2.font.size = Pt(12)
    doc.add_paragraph()

    prem = (
        f"{s.company_full_name}, именуемое в дальнейшем «Исполнитель», "
        f"в лице {s.company_director_position} {s.company_director_name}, "
        f"действующего на основании свидетельства о государственной регистрации, "
        f"с одной стороны, и\n\n"
        f"{req['full_name']}, именуемое в дальнейшем «Заказчик», "
        f"в лице {req['director_position']} {req['director_name']}, "
        f"действующего на основании Устава, с другой стороны,\n\n"
        f"совместно именуемые «Стороны», заключили настоящий Договор о нижеследующем:"
    )
    _add_paragraph(doc, prem)
    doc.add_paragraph()

    def sec(title: str, paras: list[str]) -> None:
        _section_heading(doc, title)
        _add_justify_lines(doc, paras)
        doc.add_paragraph()

    sec("1. ПРЕДМЕТ И ОБЩИЕ УСЛОВИЯ ДОГОВОРА", [
        f"1.1.\u2002Заказчик поручает, а Исполнитель принимает на себя обязательство "
        f"по выполнению рабочей документации на узел коммерческого учёта тепловой "
        f"энергии (далее — УКУТЭ), теплоносителя (далее — Проект) по объекту: "
        f"{ctx.object_address}.",
        "1.2.\u2002Проект разрабатывается в соответствии с требованиями "
        "Федерального закона от 23.11.2009 № 261-ФЗ, Постановления Правительства РФ "
        "от 18.11.2013 № 1034 (далее — Правила № 1034), Приказа Минстроя России "
        "от 17.03.2014 № 99/пр (далее — Методика), технических условий (ТУ) РСО "
        "(Приложение № 2), технической документации производителей приборов учёта.",
        "1.3.\u2002Состав Проекта приведён в Приложении № 1 к настоящему Договору.",
        "1.4.\u2002Первичные почтовые расходы, связанные с доставкой Проекта "
        "Заказчику в электронном виде, несёт Исполнитель.",
        "1.5.\u2002Настоящий Договор составлен в электронной форме и имеет юридическую "
        "силу согласно п. 2 ст. 434 ГК РФ, пунктам 10.1–10.3 настоящего Договора.",
    ])

    sec("2. ПОРЯДОК ЭЛЕКТРОННОГО ДОКУМЕНТООБОРОТА И ЮРИДИЧЕСКАЯ СИЛА ПЕРЕПИСКИ", [
        f"2.1.\u2002Стороны признают юридическую силу электронных писем с адресов: "
        f"Исполнитель: {executor_email}; Заказчик: {ce}.",
        "2.2.\u2002Доступ к электронной почте каждая Сторона осуществляет по паролю "
        "и обязуется сохранять конфиденциальность пароля.",
        "2.3.\u2002Стороны договорились о возможности использования аналога "
        "собственноручной подписи для подписания счетов и актов на основании "
        "п. 2 ст. 160 ГК РФ.",
        "2.4.\u2002Договор, приложения и акты могут быть подписаны обменом "
        "сканированными копиями по электронной почте с указанных в п. 2.1 адресов.",
        "2.5.\u2002Заказчик вместе с подписанным Договором направляет ответным письмом "
        "документы и сведения согласно таблице:",
    ])

    tbl = doc.add_table(rows=8, cols=3)
    tbl.style = "Table Grid"
    hdr = ["№", "Документ / сведения", "Обязательность"]
    for i, h in enumerate(hdr):
        cell = tbl.rows[0].cells[i]
        cell.paragraphs[0].clear()
        r = cell.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(10)
    rows = [
        ("1", "Технические условия РСО (Приложение № 2)", "обязательно"),
        ("2", "Акт разграничения балансовой принадлежности (существующие объекты)", "обязательно"),
        ("3", "План теплового пункта с местом установки УКУТЭ", "обязательно"),
        ("4", "Принципиальная схема теплового пункта", "обязательно"),
        ("5", "План подключения потребителя к тепловой сети", "обязательно"),
        ("6", "Карточка предприятия Заказчика", "обязательно"),
        ("7", "Контактные данные ответственного лица", "обязательно"),
    ]
    for ri, (a, b, c) in enumerate(rows, start=1):
        for ci, val in enumerate((a, b, c)):
            cell = tbl.rows[ri].cells[ci]
            cell.paragraphs[0].clear()
            r = cell.paragraphs[0].add_run(val)
            r.font.size = Pt(10)

    doc.add_paragraph()
    _add_justify_lines(doc, [
        f"2.6.\u2002Комплект по п. 2.5 считается переданным с поступления письма "
        f"на {executor_email}. Исполнитель вправе запросить недостающие документы; "
        f"срок по п. 5.1 исчисляется после получения всех сведений.",
    ])

    sec("3. СРОК ДЕЙСТВИЯ ДОГОВОРА", [
        f"3.1.\u2002Срок действия Договора — с даты подписания по 31.12.{now.year} г., "
        f"в части расчётов — до полного исполнения обязательств.",
        "3.2.\u2002Договор вступает в силу с даты подписания Сторонами в порядке раздела 2.",
    ])

    sec("4. СТОИМОСТЬ РАБОТ И ПОРЯДОК РАСЧЁТОВ", [
        f"4.1.\u2002Стоимость Проекта составляет {_fmt(ctx.payment_amount)}\u202fруб. "
        f"({words_total}) без НДС (УСН, ст.\u202f346.11 НК\u202fРФ). "
        f"Счёт-фактура не предоставляется.",
        f"4.2.\u2002Оплата: а) аванс {_fmt(ctx.advance_amount)}\u202fруб. ({words_adv}) "
        f"(50\u202f%) — в течение 3 банковских дней с подписания Договора; "
        f"б) остаток {_fmt(ctx.final_amount)}\u202fруб. ({words_fin}) — "
        f"в порядке раздела 5.",
        "4.3.\u2002Оплата безналичным переводом на счёт Исполнителя (раздел 13).",
        "4.4.\u2002Проценты по ст.\u202f317.1 ГК РФ Сторонами на отношения по Договору "
        "не начисляются.",
    ])

    sec("5. ПОРЯДОК ВЫПОЛНЕНИЯ, СДАЧИ И ПРИЁМКИ РАБОТ", [
        "5.1.\u2002Исполнитель выполняет Проект не более чем за 3 рабочих дня с даты, "
        "наступившей позже: а) поступления аванса; б) получения полного комплекта по п.\u202f2.5.",
        f"5.2.\u2002Проект направляется на e-mail: {ce}.",
        "5.3.\u2002Дата направления по e-mail — дата сдачи работ (ст.\u202f720 ГК РФ).",
        f"5.4.\u2002В течение 5 рабочих дней Заказчик направляет на {executor_email} "
        f"копию сопроводительного с отметкой РСО либо производит окончательный расчёт "
        f"и направляет Акт.",
        "5.5.\u2002Если в течение 3 рабочих дней нет мотивированного отказа ни Акта — "
        "Проект принят (п.\u202f4 ст.\u202f753 ГК РФ).",
        "5.6.\u2002В течение 15 рабочих дней с п.\u202f5.4 Заказчик производит окончательный "
        "расчёт либо направляет замечания РСО на Листе согласования (Приложение № 3).",
        "5.7.\u2002Если замечаний и оплаты нет — Проект считается согласованным (ст.\u202f421 ГК РФ).",
        "5.8.\u2002Замечания РСО устраняются за счёт Исполнителя в срок 5 рабочих дней, "
        "если недостатки по его вине; иные случаи — по доп. соглашению.",
    ])

    sec("6. ПРАВА И ОБЯЗАННОСТИ СТОРОН", [
        "6.1.1.\u2002Своевременно и в полном объёме производить оплату работ "
        "в соответствии с разделом 4 Договора.",
        "6.1.2.\u2002Предоставить Исполнителю полный комплект документов по п.\u202f2.5. "
        "Заказчик несёт ответственность за полноту и достоверность информации.",
        "6.1.3.\u2002Своевременно производить приёмку Проекта в соответствии с п.\u202f5.4.",
        "6.1.4.\u2002Обеспечить передачу Проекта в РСО и направить Исполнителю копию "
        "сопроводительного письма согласно п.\u202f5.4.",
        "6.2.1.\u2002Запрашивать у Исполнителя информацию о ходе выполнения Договора.",
        "6.2.2.\u2002Направлять мотивированный отказ от приёмки Проекта в порядке п.\u202f5.5.",
        "6.3.1–6.3.5.\u2002Исполнитель выполняет Проект в срок п.\u202f5.1, устраняет замечания "
        "РСО по п.\u202f5.8, информирует Заказчика, не передаёт Проект третьим лицам "
        "без согласования до подписания Акта.",
        "6.4.1–6.4.3.\u2002Исполнитель вправе привлекать субподрядчиков, использовать "
        "материалы в портфолио (без реквизитов Заказчика), включать приборы из ФИФ по метрологии.",
        "6.5.\u2002Исполнитель несёт финансовый риск перед Заказчиком только в пределах "
        "суммы настоящего Договора.",
    ])

    sec("7. ОТВЕТСТВЕННОСТЬ СТОРОН", [
        "7.1.\u2002За неисполнение или ненадлежащее исполнение обязательств Стороны несут "
        "ответственность в соответствии с законодательством РФ, кроме случаев, "
        "специально оговоренных Договором.",
        "7.2.\u2002За просрочку оплаты по пп.\u202f4.2.а и 4.2.б Исполнитель вправе "
        "требовать пени 1\u202f% от неуплаченной суммы за каждый день просрочки, "
        "но не более суммы Договора.",
        "7.3.\u2002При просрочке предоставления документов по п.\u202f2.5 более чем на "
        "15 календарных дней с даты оплаты аванса Исполнитель вправе приостановить "
        "работы и/или расторгнуть Договор с удержанием аванса.",
        "7.4.\u2002Если Заказчик не исполнил п.\u202f5.4 в течение 5 календарных дней "
        "с даты направления Проекта, Проект считается принятым, окончательный расчёт — "
        "в безусловном порядке.",
        "7.5–7.6.\u2002Споры разрешаются переговорами; претензионный порядок обязателен, "
        "ответ на претензию — в течение 10 рабочих дней.",
        "7.7.\u2002При недостижении согласия спор передаётся в Арбитражный суд Амурской области.",
        "7.8.\u2002Если Заказчик принял работу без проверки, он лишается права ссылаться "
        "на недостатки, которые могли быть выявлены при обычном приёме (п.\u202f3 ст.\u202f720 ГК РФ).",
    ])

    sec("8. ИЗМЕНЕНИЕ И РАСТОРЖЕНИЕ ДОГОВОРА", [
        "8.1.\u2002Изменение и расторжение по соглашению Сторон.",
        "8.2.\u2002Дополнения действительны в письменной форме (раздел 2).",
        "8.3.\u2002Одностороннее расторжение с уведомлением за 2 рабочих дня при "
        "существенных нарушениях.",
    ])

    sec("9. ОСОБЫЕ УСЛОВИЯ", [
        "9.1.\u2002Заказчик не вправе отказаться от согласования при соответствии "
        "Приложениям № 1 и 2.",
        "9.2.\u2002Исключительное право на Проект переходит к Заказчику после полной оплаты.",
    ])

    sec("10. ЭЛЕКТРОННАЯ ФОРМА ДОГОВОРА", [
        "10.1–10.3.\u2002Обмен по e-mail как письменная форма (ст.\u202f434 ГК РФ, ФЗ № 63).",
    ])

    sec("11. ОБСТОЯТЕЛЬСТВА НЕПРЕОДОЛИМОЙ СИЛЫ", [
        "11.1–11.3.\u2002Освобождение от ответственности при форс-мажоре, уведомление за 5 "
        "рабочих дней, продление сроков.",
    ])

    sec("12. КОНФИДЕНЦИАЛЬНОСТЬ", [
        "12.1.\u2002Стороны обязуются сохранять конфиденциальность информации по Договору.",
    ])

    sec("13. АНТИКОРРУПЦИОННАЯ ОГОВОРКА", [
        "13.1–13.5.\u2002Стороны подтверждают легитимность деятельности, запрет схем "
        "коррупции, право приостановки и расторжения при нарушениях.",
    ])

    sec("14. ПРИЛОЖЕНИЯ", [
        "14.1.\u2002Неотъемлемыми частями являются: Приложение № 1 — состав документации; "
        "Приложение № 2 — ТУ РСО; Приложение № 3 — Лист согласования.",
    ])

    _section_heading(doc, "15. РЕКВИЗИТЫ И ПОДПИСИ СТОРОН")
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(8)
    table.columns[1].width = Cm(8)
    _set_table_borders(table, inner_v=True)
    left_cell, right_cell = table.rows[0].cells
    cl_phone = req.get("phone") or "—"
    cl_ogrn = req.get("ogrn") or "—"
    _cell_lines(left_cell, [
        ("Исполнитель", True),
        "",
        s.company_full_name,
        f"ИНН\u202f{s.company_inn}",
        "КПП\u202f—",
        f"ОГРН\u202f{s.company_ogrn}",
        s.company_address or "—",
        f"Тел.:\u202f—",
        f"E-mail:\u202f{executor_email}",
        f"р/с\u202f{s.company_settlement_account}",
        f"в {s.company_bank_name}",
        f"БИК\u202f{s.company_bik}",
        f"к/с\u202f{s.company_corr_account}",
        "",
        f"__________\u202f/\u202f{s.company_director_name}",
        "М.П.",
    ])
    _cell_lines(right_cell, [
        ("Заказчик", True),
        "",
        req["full_name"],
        f"ИНН\u202f{req['inn']}",
        f"КПП\u202f{client_kpp}",
        f"ОГРН\u202f{cl_ogrn}",
        req["legal_address"],
        f"Тел.:\u202f{cl_phone}",
        f"E-mail:\u202f{ce}",
        f"р/с\u202f{req['settlement_account']}",
        f"в {req['bank_name']}",
        f"БИК\u202f{req['bik']}",
        f"к/с\u202f{req['corr_account']}",
        "",
        f"__________\u202f/\u202f{req['director_name']}",
        "М.П.",
    ])
    doc.add_paragraph()


def _build_appendix_1(doc: Document, ctx: _ContractContext) -> None:
    """Приложение № 1 — состав рабочей документации УКУТЭ."""
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ПРИЛОЖЕНИЕ № 1")
    r.bold = True
    r.font.size = Pt(14)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(
        f"к Договору № {ctx.contract_number} от «___» __________ {datetime.now().year} г."
    )
    r2.font.size = Pt(12)
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("Состав рабочей документации УКУТЭ")
    r3.bold = True
    r3.font.size = Pt(12)
    doc.add_paragraph()
    _add_paragraph(
        doc,
        "Исполнитель разрабатывает и передаёт Заказчику следующий комплект в электронном виде (PDF):",
    )
    items = [
        "1. Пояснительная записка.",
        "2. Общие данные.",
        "3. План подключения потребителя к тепловой сети.",
        "4. Принципиальная схема теплового пункта с узлом учёта.",
        "5. План теплового пункта с указанием мест установки датчиков и приборов учёта.",
        "6. Электрические и монтажные схемы подключения приборов учёта.",
        "7. Настроечная база данных для теплоэнергоконтроллера (вычислителя).",
        "8. Схема пломбирования средств измерений.",
        "9. Форма отчётной ведомости показаний приборов учёта.",
        "10. Монтажные схемы установки расходомера и датчиков.",
        "11. Спецификация оборудования и материалов.",
        "12. Расчёт потерь напора теплоносителя.",
    ]
    for line in items:
        _add_paragraph(doc, line)
    doc.add_paragraph()
    _appendix_signatures_row(doc)


def _build_appendix_2_header(doc: Document, ctx: _ContractContext) -> None:
    """Заголовок Приложения № 2 и вводный блок без текста о количестве листов / без растра."""
    doc.add_page_break()
    for title, size, bold in [
        ("ПРИЛОЖЕНИЕ № 2", 14, True),
        (
            f"к Договору № {ctx.contract_number} от «___» __________ "
            f"{datetime.now().year} г.",
            12,
            False,
        ),
        ("Технические условия РСО", 12, True),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(title)
        r.bold = bold
        r.font.size = Pt(size)
    doc.add_paragraph()
    rn = ctx.rso_name or "—"
    tnum = ctx.tu_number or "—"
    td = ctx.tu_date or "—"
    tv = ctx.tu_valid_to or "—"
    _add_paragraph(
        doc,
        f"Технические условия на проектирование узла учёта тепловой энергии, "
        f"теплоносителя, выданные теплоснабжающей организацией {rn} № {tnum} от {td}, "
        f"действительные до {tv}.",
    )


def _add_signatures_block(doc: Document) -> None:
    """Подписи сторон в конце Приложения № 2 (и аналогично структурируемых блоков)."""
    doc.add_paragraph()
    _appendix_signatures_row(doc)


def _build_appendix_3(doc: Document, ctx: _ContractContext) -> None:
    """Приложение № 3 — Лист согласования."""
    doc.add_page_break()
    for title, bold, fs in [
        ("ПРИЛОЖЕНИЕ № 3", True, 14),
        (
            f"к Договору № {ctx.contract_number} от «___» __________ "
            f"{datetime.now().year} г.",
            False,
            12,
        ),
        ("Лист согласования", True, 12),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(title)
        r.bold = bold
        r.font.size = Pt(fs)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(f"«___» _____________ {datetime.now().year} г.")
    r.font.size = Pt(12)
    doc.add_paragraph()
    coord = doc.add_table(rows=11, cols=2)
    coord.style = "Table Grid"
    h1 = coord.rows[0].cells[0]
    h2 = coord.rows[0].cells[1]
    h1.paragraphs[0].clear()
    h2.paragraphs[0].clear()
    r1 = h1.paragraphs[0].add_run("Лист проекта")
    r1.bold = True
    r2 = h2.paragraphs[0].add_run("Замечание")
    r2.bold = True
    for ri in range(1, 11):
        for ci in range(2):
            coord.rows[ri].cells[ci].paragraphs[0].clear()
    doc.add_paragraph()
    _appendix_signatures_row(doc)


def _build_full_contract_document(
    ctx: _ContractContext,
    *,
    tu_png_paths: list[Path],
    use_embedded_tu: bool,
) -> Document:
    """Собирает Document: основной текст, приложения 1–3, в Приложении 2 — ТУ или заглушка.

    Args:
        ctx: реквизиты и параметры договора.
        tu_png_paths: пути к PNG-страницам ТУ (пустой список — текст-заглушка вместо скана).
        use_embedded_tu: если True и tu_png_paths не пуст — вставляем картинки и текст про листы.
    """
    global _current_req_snapshot
    _current_req_snapshot = ctx.requisites
    try:
        doc = Document()
        _apply_contract_page_settings(doc)
        _build_main_contract(doc, ctx)
        _build_appendix_1(doc, ctx)
        _build_appendix_2_header(doc, ctx)
        if use_embedded_tu and tu_png_paths:
            _add_paragraph(
                doc,
                f"ТУ приложены к настоящему Договору на {len(tu_png_paths)} "
                f"листах и являются неотъемлемой частью Договора.",
            )
            for png_path in tu_png_paths:
                doc.add_page_break()
                doc.add_picture(str(png_path), width=Cm(16.5))
        else:
            _add_paragraph(
                doc,
                "ТУ прилагаются отдельным файлом в электронном виде "
                "и являются неотъемлемой частью Договора.",
            )
        _add_signatures_block(doc)
        _build_appendix_3(doc, ctx)
        return doc
    finally:
        _current_req_snapshot = None


def _save_contract_docx(doc: Document, order_id_short: str) -> Path:
    tmp = tempfile.NamedTemporaryFile(
        suffix=".docx",
        prefix=f"dogovor_{order_id_short}_",
        delete=False,
        dir="/tmp",
    )
    tmp.close()
    doc.save(tmp.name)
    return Path(tmp.name)


def _build_contract_without_tu(ctx: _ContractContext) -> Path:
    """Аварийная генерация договора без вставки страниц ТУ (только заглушка в Приложении 2)."""
    logger.info(
        "Сборка договора order=%s без встроенных страниц ТУ (аварийный режим)",
        ctx.order_id_short,
    )
    doc = _build_full_contract_document(
        ctx,
        tu_png_paths=[],
        use_embedded_tu=False,
    )
    return _save_contract_docx(doc, ctx.order_id_short)


# ═══════════════════════════════════════════════════════════════════════════════
# Публичные функции
# ═══════════════════════════════════════════════════════════════════════════════


def generate_contract_number(order_id: str) -> str:
    """Формирует номер договора: UUTE-YYYYMMDD-xxxx.

    Пример: «UUTE-20260415-a1b2»
    """
    return f"UUTE-{datetime.now().strftime('%Y%m%d')}-{order_id[:4]}"


def generate_contract(
    order_id_short: str,
    contract_number: str,
    object_address: str,
    client_name: str,
    payment_amount: int,
    advance_amount: int,
    requisites: dict,
    client_email: str | None = None,
    tu_file_path: Path | None = None,
    rso_name: str | None = None,
    tu_number: str | None = None,
    tu_date: str | None = None,
    tu_valid_to: str | None = None,
) -> Path:
    """Генерирует DOCX договора УКУТЭ по шаблону с приложениями 1–3.

    Приложение №2: при переданном ``tu_file_path`` страницы PDF ТУ раструются
    в PNG (PyMuPDF) и вставляются с шириной 16.5 см. Если итоговый DOCX превышает
    ~25 МБ, DPI снижается по лестнице 150 → 120 → 100. Если даже на минимальном
    DPI размер велик — генерируется договор без встроенных страниц ТУ, с
    заглушкой и записью уровня ERROR в лог.

    Args:
        order_id_short: первые 8 символов UUID заявки (префикс имени файла).
        contract_number: номер договора из ``generate_contract_number``.
        object_address: адрес объекта теплоснабжения.
        client_name: ФИО/наименование (запасной контакт).
        payment_amount: полная стоимость, руб.
        advance_amount: аванс 50 %, руб.
        requisites: реквизиты клиента (``CompanyRequisites.model_dump()``).
        client_email: e-mail клиента.
        tu_file_path: абсолютный путь к PDF ТУ на диске (или ``None``).
        rso_name, tu_number, tu_date, tu_valid_to: поля для шапки Приложения 2
            из ``parsed_params`` (``rso`` / ``document``); могут быть ``None``.

    Returns:
        Path к временному ``.docx`` в ``/tmp``. Вызывающий код удаляет файл
        после отправки. Временные PNG после выхода из функции не остаются.
    """
    final_amount = payment_amount - advance_amount
    ctx = _ContractContext(
        order_id_short=order_id_short,
        contract_number=contract_number,
        object_address=object_address,
        client_name=client_name,
        payment_amount=payment_amount,
        advance_amount=advance_amount,
        final_amount=final_amount,
        requisites=requisites,
        client_email=client_email,
        rso_name=rso_name,
        tu_number=tu_number,
        tu_date=tu_date,
        tu_valid_to=tu_valid_to,
    )

    final_docx_path: Path | None = None
    tu_page_paths: list[Path] = []

    def _log_attempt(attempt_dpi: int, tmp_path: Path, n_pages: int) -> None:
        size = os.path.getsize(tmp_path)
        logger.info(
            "Договор сгенерирован: order=%s, dpi=%d, pages_tu=%d, size=%d МБ",
            order_id_short,
            attempt_dpi,
            n_pages,
            size // 1024 // 1024,
        )

    tu_ready = tu_file_path is not None and tu_file_path.exists()

    if not tu_ready:
        doc = _build_full_contract_document(
            ctx,
            tu_png_paths=[],
            use_embedded_tu=False,
        )
        tmp_path = _save_contract_docx(doc, order_id_short)
        size = os.path.getsize(tmp_path)
        logger.info(
            "Договор сгенерирован без вложения ТУ: order=%s, pages_tu=0, size=%d МБ",
            order_id_short,
            size // 1024 // 1024,
        )
        return tmp_path

    for attempt_dpi in _TU_DPI_LADDER:
        _cleanup_png_files(tu_page_paths)
        tu_page_paths = []
        tu_page_paths = _render_tu_pages_to_png(tu_file_path, dpi=attempt_dpi)
        if not tu_page_paths:
            logger.warning(
                "Страницы ТУ не получены (order=%s, dpi=%s), вставляем заглушку",
                order_id_short,
                attempt_dpi,
            )
            doc = _build_full_contract_document(
                ctx,
                tu_png_paths=[],
                use_embedded_tu=False,
            )
            tmp_path = _save_contract_docx(doc, order_id_short)
            _log_attempt(attempt_dpi, tmp_path, 0)
            return tmp_path

        doc = _build_full_contract_document(
            ctx,
            tu_png_paths=tu_page_paths,
            use_embedded_tu=True,
        )
        tmp_path = _save_contract_docx(doc, order_id_short)
        size = os.path.getsize(tmp_path)
        _log_attempt(attempt_dpi, tmp_path, len(tu_page_paths))

        if size <= _MAX_DOCX_SIZE_BYTES:
            final_docx_path = tmp_path
            break

        logger.warning(
            "Договор %s МБ > лимита на DPI=%d, понижаем...",
            size // 1024 // 1024,
            attempt_dpi,
        )
        Path(tmp_path).unlink(missing_ok=True)
    else:
        logger.error(
            "Договор order=%s не умещается в %d МБ даже на DPI=%d. "
            "ТУ слишком тяжёлый. Нужно ручное вмешательство.",
            order_id_short,
            _MAX_DOCX_SIZE_BYTES // 1024 // 1024,
            _TU_DPI_LADDER[-1],
        )
        final_docx_path = _build_contract_without_tu(ctx)

    _cleanup_png_files(tu_page_paths)
    assert final_docx_path is not None
    return final_docx_path


def generate_invoice(
    order_id_short: str,
    contract_number: str,
    object_address: str,
    payment_amount: int,
    advance_amount: int,
    requisites: dict,
    is_advance: bool = True,
) -> Path:
    """Генерирует DOCX счёта на оплату (аванс или окончательный расчёт).

    Args:
        order_id_short: Первые 8 символов UUID заявки (для имени файла).
        contract_number: Номер договора.
        object_address:  Адрес объекта.
        payment_amount:  Полная стоимость работ, руб.
        advance_amount:  Сумма аванса (50%), руб.
        requisites:      Реквизиты клиента из CompanyRequisites.model_dump().
        is_advance:      True = счёт на 50% аванс, False = счёт на остаток 50%.

    Returns:
        Path к временному .docx файлу. Вызывающий код удаляет после отправки.
    """
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(2)
    sec.right_margin = Cm(1.5)

    s = settings
    amount = advance_amount if is_advance else (payment_amount - advance_amount)
    invoice_label = "аванс 50%" if is_advance else "окончательный расчёт 50%"
    now = datetime.now()
    client_kpp = requisites.get("kpp") or "—"
    words_amount = number_to_words_ru(amount)

    # ── Шапка: реквизиты банка получателя ────────────────────────────────────
    # Стандартный формат российского счёта: 4 строки, 2 колонки
    hdr = doc.add_table(rows=4, cols=2)
    hdr.style = "Table Grid"
    hdr.autofit = False
    hdr.columns[0].width = Cm(10)
    hdr.columns[1].width = Cm(7.5)

    def _hdr(row: int, col: int, lines: str, bold: bool = False) -> None:
        cell = hdr.rows[row].cells[col]
        cell.paragraphs[0].clear()
        _cell_add_break_lines(cell, lines)
        if bold:
            for run in cell.paragraphs[0].runs:
                run.bold = True

    _hdr(0, 0, f"Банк получателя:\n{s.company_bank_name}")
    _hdr(0, 1, f"БИК\n{s.company_bik}")
    _hdr(1, 0, "")
    _hdr(1, 1, f"Сч.\u202f№\n{s.company_corr_account}")
    _hdr(2, 0, f"ИНН\u202f{s.company_inn}")
    _hdr(2, 1, "")
    _hdr(3, 0, f"Получатель:\n{s.company_full_name}", bold=True)
    _hdr(3, 1, f"Сч.\u202f№\n{s.company_settlement_account}")

    doc.add_paragraph()

    # ── Заголовок счёта ───────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        f"СЧЁТ НА ОПЛАТУ\u202f№\u202f{contract_number}\n"
        f"от {_ru_date(now)}"
    )
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()

    # ── Поставщик и покупатель ────────────────────────────────────────────────
    def _info_line(label: str, value: str) -> None:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r1 = p.add_run(f"{label}:\u2003")
        r1.bold = True
        r1.font.size = Pt(11)
        r2 = p.add_run(value)
        r2.font.size = Pt(11)

    _info_line(
        "Поставщик (Исполнитель)",
        f"{s.company_full_name}, ИНН\u202f{s.company_inn}, "
        f"{s.company_address or '—'}",
    )
    _info_line(
        "Покупатель (Заказчик)",
        f"{requisites['full_name']}, ИНН\u202f{requisites['inn']}, "
        f"КПП\u202f{client_kpp}, {requisites['legal_address']}",
    )
    _info_line("Основание", f"Договор\u202f№\u202f{contract_number}")

    doc.add_paragraph()

    # ── Таблица услуг ─────────────────────────────────────────────────────────
    # Ширины: 1.0 + 8.5 + 1.5 + 1.5 + 2.5 + 2.5 = 17.5 см
    col_headers = ["№", "Наименование работ (услуг)", "Кол.", "Ед.", "Цена, руб.", "Сумма, руб."]
    col_widths = [Cm(1.0), Cm(8.5), Cm(1.5), Cm(1.5), Cm(2.5), Cm(2.5)]

    svc_tbl = doc.add_table(rows=2, cols=6)
    svc_tbl.style = "Table Grid"
    svc_tbl.autofit = False

    # Заголовок таблицы
    for i, (header, width) in enumerate(zip(col_headers, col_widths)):
        cell = svc_tbl.rows[0].cells[i]
        cell.width = width
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(header)
        run.bold = True
        run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Строка с услугой
    service_name = (
        f"Разработка проекта УУТЭ по объекту:\n{object_address} ({invoice_label})"
    )
    row_data = [
        ("1",              WD_ALIGN_PARAGRAPH.CENTER),
        (service_name,     WD_ALIGN_PARAGRAPH.LEFT),
        ("1",              WD_ALIGN_PARAGRAPH.CENTER),
        ("усл.",           WD_ALIGN_PARAGRAPH.CENTER),
        (_fmt(amount),     WD_ALIGN_PARAGRAPH.RIGHT),
        (_fmt(amount),     WD_ALIGN_PARAGRAPH.RIGHT),
    ]
    for i, (text, align) in enumerate(row_data):
        cell = svc_tbl.rows[1].cells[i]
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(text)
        run.font.size = Pt(10)
        cell.paragraphs[0].alignment = align

    doc.add_paragraph()

    # ── Итого ─────────────────────────────────────────────────────────────────
    for text, bold in [
        (f"Итого:\u2003{_fmt(amount)}\u202fруб.", True),
        ("В том числе НДС: не облагается (УСН, ст.\u202f346.11 НК\u202fРФ).", False),
        (
            f"Всего к оплате:\u2003{_fmt(amount)}\u202fруб. "
            f"({words_amount[0].upper() + words_amount[1:]}).",
            True,
        ),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph()

    # ── Подпись ───────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(
        f"{s.company_director_position}:\u2003"
        f"__________\u202f/\u202f{s.company_director_name}"
    )
    run.font.size = Pt(12)

    # ── Сохранить во временный файл ───────────────────────────────────────────
    tmp = tempfile.NamedTemporaryFile(
        suffix=".docx",
        prefix=f"schet_{order_id_short}_",
        delete=False,
    )
    tmp.close()
    doc.save(tmp.name)
    return Path(tmp.name)
