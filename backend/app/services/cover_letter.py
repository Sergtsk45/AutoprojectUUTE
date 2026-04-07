"""Генератор сопроводительного письма в РСО (.docx).

Создаёт предзаполненный шаблон письма для клиента:
  - отправитель: заявитель (данные из ТУ — applicant.*)
  - получатель: РСО (данные из ТУ — rso.*)
  - поля исходящего номера/даты — пустые, клиент заполняет вручную перед подачей
"""

import tempfile
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.services.tu_schema import TUParsedData


def generate_cover_letter(parsed: TUParsedData, order_id_short: str) -> Path:
    """Генерирует DOCX сопроводительного письма в РСО.

    Args:
        parsed: Извлечённые данные из ТУ (applicant.*, rso.*, document.*, object.*).
        order_id_short: Первые 8 символов UUID заявки — используется в имени файла.

    Returns:
        Path к временному .docx файлу. Вызывающий код обязан удалить файл после отправки.
    """
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)

    applicant_name = parsed.applicant.applicant_name or "________________________"
    applicant_address = parsed.applicant.applicant_address or "________________________"
    contact_person = parsed.applicant.contact_person or "________________________"
    rso_name = parsed.rso.rso_name or "________________________"
    rso_address = parsed.rso.rso_address or "________________________"
    tu_number = parsed.document.tu_number or "___"
    tu_date = parsed.document.tu_date or "________"
    object_address = parsed.object.object_address or "________________________"

    # ── Кому ─────────────────────────────────────────────────────────────
    to_para = doc.add_paragraph()
    to_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = to_para.add_run(f"{rso_name}\n{rso_address}")
    run.font.size = Pt(12)

    doc.add_paragraph()

    # ── От кого ───────────────────────────────────────────────────────────
    from_para = doc.add_paragraph()
    from_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = from_para.add_run(f"От: {applicant_name}\n{applicant_address}")
    run.font.size = Pt(12)

    doc.add_paragraph()

    # ── Исходящий номер и дата ────────────────────────────────────────────
    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = meta_para.add_run("Исх. № __________ от «____» ____________ 20___ г.")
    run.font.size = Pt(12)

    doc.add_paragraph()

    # ── Заголовок письма ──────────────────────────────────────────────────
    subj_para = doc.add_paragraph()
    subj_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subj_para.add_run(
        "О направлении проекта узла учёта тепловой энергии (УУТЭ) на согласование"
    )
    run.bold = True
    run.font.size = Pt(13)

    doc.add_paragraph()

    # ── Тело письма ───────────────────────────────────────────────────────
    body_text = (
        f"В соответствии с Приказом Минстроя России №\u202f1036/пр "
        f"«Правила коммерческого учёта тепловой энергии, теплоносителя», "
        f"а также техническими условиями № {tu_number} от {tu_date}, "
        f"выданными {rso_name}, направляем на согласование проект "
        f"узла учёта тепловой энергии (УУТЭ) по объекту:\n\n"
        f"{object_address}.\n\n"
        f"Просим рассмотреть представленный проект и согласовать его "
        f"в установленные сроки в соответствии с действующим законодательством.\n\n"
        f"По вопросам согласования просим обращаться: {contact_person}."
    )

    body_para = doc.add_paragraph()
    body_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = body_para.add_run(body_text)
    run.font.size = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph()

    # ── Подпись ───────────────────────────────────────────────────────────
    sign_para = doc.add_paragraph()
    sign_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = sign_para.add_run(
        f"{contact_person or applicant_name}\n\n"
        "Подпись: ______________________\n\n"
        "М.П."
    )
    run.font.size = Pt(12)

    # ── Сохранить во временный файл ───────────────────────────────────────
    tmp = tempfile.NamedTemporaryFile(
        suffix=".docx",
        prefix=f"soprovod_{order_id_short}_",
        delete=False,
    )
    tmp.close()
    doc.save(tmp.name)

    return Path(tmp.name)
